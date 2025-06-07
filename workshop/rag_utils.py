from abc import ABC, abstractmethod
import os
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup
import docx
import json
import re
import math
import base64
from typing import Any, List, Dict, Optional, Tuple, Union
from dataclasses import dataclass
from langchain_core.messages import HumanMessage
from settings.llm_api_aggregator import WWSettingsManager, WWApiAggregator
import tiktoken
import fitz
import pymupdf4llm
from PyQt5.QtCore import QThread, pyqtSignal, Qt, QPoint
from PyQt5.QtWidgets import QDialog, QVBoxLayout, QLineEdit, QPushButton, QLabel, QListWidget, QMessageBox, QMenu, QAction, QInputDialog, QTextEdit, QHBoxLayout, QCheckBox, QListWidgetItem, QWidget, QShortcut
from PyQt5.QtGui import QTextDocument, QTextCursor, QKeySequence

from abc import ABC, abstractmethod
import os
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup
import docx

# Abstract base class for all document processors
class DocumentProcessor(ABC):
    @abstractmethod
    def load_document(self, file_path: str) -> Tuple[int, Optional[str]]:
        # Returns the number of sections (pages, chapters, etc.) and an optional error message
        pass
    
    @abstractmethod
    def convert_to_markdown(self, file_path: str, pages_or_sections: List[int]) -> Tuple[str, Optional[str]]:
        # Converts specified sections to markdown, returns markdown text and an optional error message
        pass

# EPUB processor using ebooklib and BeautifulSoup
class EpubProcessor(DocumentProcessor):
    @staticmethod
    def load_document(epub_path: str) -> Tuple[int, Optional[str]]:
        # Load EPUB and count chapters as sections
        try:
            book = epub.read_epub(epub_path)
            chapters = list(book.get_items_of_type(ebooklib.ITEM_DOCUMENT))
            return len(chapters), None
        except Exception as e:
            return 0, f"Error loading EPUB: {e}"
    
    @staticmethod
    def convert_to_markdown(epub_path: str, chapters: List[int] = None) -> Tuple[str, Optional[str]]:
        """
        For each chapter indicated:
        - parses the HTML
        - extracts a list of <p>...</p>.
        - gives each paragraph a separate entry
        """
        try:
            book = epub.read_epub(epub_path)
            all_items = list(book.get_items_of_type(ebooklib.ITEM_DOCUMENT))
            # wybieramy rozdziały
            if chapters is not None:
                selected = [all_items[i] for i in chapters if 0 <= i < len(all_items)]
            else:
                selected = all_items
            
            markdown_sections: List[str] = []
            for chap in selected:
                raw = chap.get_content().decode('utf-8')
                soup = BeautifulSoup(raw, 'html.parser')
                paras = soup.find_all('p')
                for p in paras:
                    text = p.get_text(strip=True)
                    if text:
                        markdown_sections.append(text)
                # optional: insert separator between chapters
                markdown_sections.append('---')
            
            # remove the last separator, if any
            if markdown_sections and markdown_sections[-1] == '---':
                markdown_sections.pop()
            
            return '\n\n'.join(markdown_sections), None
        except Exception as e:
            return "", f"Error converting EPUB: {e}"

# DOCX processor using python-docx
class DocxProcessor(DocumentProcessor):
    @staticmethod
    def load_document(docx_path: str) -> Tuple[int, Optional[str]]:
        # Load DOCX and count paragraphs as sections
        try:
            doc = docx.Document(docx_path)
            return len(doc.paragraphs), None
        except Exception as e:
            return 0, f"Error loading DOCX: {e}"
    
    @staticmethod
    def convert_to_markdown(docx_path: str, paragraphs: List[int] = None) -> Tuple[str, Optional[str]]:
        # Convert specified paragraphs (or all if None) to markdown
        try:
            doc = docx.Document(docx_path)
            all_paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            if paragraphs is not None:
                selected_paragraphs = [all_paragraphs[i] for i in paragraphs if i < len(all_paragraphs)]
            else:
                selected_paragraphs = all_paragraphs
            
            return '\n\n'.join(selected_paragraphs), None
        except Exception as e:
            return "", f"Error converting DOCX: {e}"

# TXT/Markdown processor for plain text files
class TextProcessor(DocumentProcessor):
    @staticmethod
    def load_document(text_path: str) -> Tuple[int, Optional[str]]:
        # Load text file and count lines as sections
        try:
            with open(text_path, 'r', encoding='utf-8') as f:
                content = f.read()
            lines = content.splitlines()
            return len(lines), None
        except Exception as e:
            return 0, f"Error loading text file: {e}"
    
    @staticmethod
    def convert_to_markdown(text_path: str, lines: List[int] = None) -> Tuple[str, Optional[str]]:
        # Convert specified lines (or all if None) to markdown (essentially pass-through)
        try:
            with open(text_path, 'r', encoding='utf-8') as f:
                all_lines = f.readlines()
            
            if lines is not None:
                selected_lines = [all_lines[i] for i in lines if i < len(all_lines)]
            else:
                selected_lines = all_lines
            
            return ''.join(selected_lines), None
        except Exception as e:
            return "", f"Error converting text file: {e}"

# HTML processor using BeautifulSoup
class HtmlProcessor(DocumentProcessor):
    @staticmethod
    def load_document(html_path: str) -> Tuple[int, Optional[str]]:
        # Load HTML and count significant sections (p, div, section tags) as sections
        try:
            with open(html_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f, 'html.parser')
            return len(soup.find_all(['p', 'div', 'section'])), None
        except Exception as e:
            return 0, f"Error loading HTML: {e}"
    
    @staticmethod
    def convert_to_markdown(html_path: str, sections: List[int] = None) -> Tuple[str, Optional[str]]:
        # Convert specified sections (or all if None) to markdown
        try:
            with open(html_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f, 'html.parser')
            all_sections = soup.find_all(['p', 'div', 'section'])
            
            if sections is not None:
                selected_sections = [all_sections[i] for i in sections if i < len(all_sections)]
            else:
                selected_sections = all_sections
            
            markdown_content = []
            for section in selected_sections:
                text = section.get_text(separator='\n\n')
                markdown_content.append(text)
            
            return '\n\n---\n\n'.join(markdown_content), None
        except Exception as e:
            return "", f"Error converting HTML: {e}"

# Factory to instantiate the correct processor based on file extension
class DocumentProcessorFactory:
    @staticmethod
    def get_processor(file_path: str) -> DocumentProcessor:
        # Determine file extension and return appropriate processor
        extension = os.path.splitext(file_path)[1].lower()
        
        if extension == '.pdf':
            return PdfProcessor()
        elif extension == '.epub':
            return EpubProcessor()
        elif extension == '.docx':
            return DocxProcessor()
        elif extension in ['.txt', '.md']:
            return TextProcessor()
        elif extension == '.html':
            return HtmlProcessor()
        else:
            raise ValueError(f"Unsupported file format: {extension}")

@dataclass
class AppSettings:
    last_pdf_path_manual: str = ""
    last_pdf_path_qa: str = ""
    last_from_page_manual: int = 0
    last_to_page_manual: int = 0
    last_chunk_size: int = 20000
    default_prompt: str = ""
    
class VisionMessage(HumanMessage):
    """Custom Message class for vision-based LLMs"""
    def __init__(self, content: Union[str, List[Dict[str, Any]]]):
        super().__init__(content=content)

class TokenCounter:
    @staticmethod
    def count_tokens(text: str, encoding_name: str = 'cl100k_base') -> int:
        encoder = tiktoken.get_encoding(encoding_name)
        return len(encoder.encode(text))

    @staticmethod
    def get_encoder(encoding_name: str = 'cl100k_base'):
        return tiktoken.get_encoding(encoding_name)

class PdfProcessor(DocumentProcessor):
    ABBREVIATIONS = {'np', 'dr', 'mgr', 'itp', 'e.g', 'i.e', 'etc'}
    SENTENCE_SPLIT_REGEX = re.compile(r'(?<=[\.\!\?])\s+')
    STRUCTURE_REGEX = re.compile(r'^(#{1,6}\s+|```|\|)', re.MULTILINE)
    PARAGRAPH_SPLIT_REGEX = re.compile(r'\n{2,}')

    @staticmethod
    def load_document(pdf_path: str) -> Tuple[int, Optional[str]]:
        # Load PDF and count pages (unchanged implementation)
        try:
            doc = fitz.open(pdf_path)
            page_count = doc.page_count
            doc.close()
            return page_count, None
        except Exception as e:
            return 0, f"Error loading PDF: {e}"

    @staticmethod
    def convert_to_markdown(pdf_path: str, pages: List[int]) -> Tuple[str, Optional[str]]:
        # Convert specified pages to markdown (unchanged implementation)
        try:
            markdown_text = pymupdf4llm.to_markdown(pdf_path, pages=pages)
            if not markdown_text.strip():
                return "", "No extractable text in PDF."
            return markdown_text, None
        except Exception as e:
            return "", f"Error converting PDF: {e}"

    @staticmethod
    def preprocess(text: str) -> str:
        # Preprocess text by removing hyphenated line breaks
        return re.sub(r'-\n\s*', '', text)

    @classmethod
    def split_paragraphs(cls, text: str) -> List[str]:
        # Split text into paragraphs
        parts = re.split(r'\n{2,}|\.(?=\s+[A-Z])|(?<=[.!?])\s+(?=[A-Z])', text)
        paragraphs: List[str] = []
        current_paragraph = ""

        for part in parts:
            part = part.strip()
            if not part:
                continue

            if current_paragraph and not current_paragraph[-1].isalpha():
                current_paragraph += " " + part
            else:
                if paragraphs and paragraphs[-1] == "":
                    paragraphs.pop()
                paragraphs.append(current_paragraph)
                current_paragraph = part

        if current_paragraph:
            paragraphs.append(current_paragraph)

        return paragraphs

    @classmethod
    def is_structural(cls, text: str) -> bool:
        # Check if text contains structural markdown elements
        for line in text.splitlines():
            if cls.STRUCTURE_REGEX.match(line):
                return True
        return False

    @classmethod
    def split_sentences(cls, text: str) -> List[str]:
        # Split text into sentences, preserving abbreviations
        parts = cls.SENTENCE_SPLIT_REGEX.split(text)
        sentences: List[str] = []
        i = 0
        while i < len(parts):
            segment = parts[i]
            lower = segment.strip().lower()
            if any(lower.endswith(abbr + '.') for abbr in cls.ABBREVIATIONS) and i + 1 < len(parts):
                segment = segment + ' ' + parts[i + 1]
                i += 2
            else:
                i += 1
            sentences.append(segment.strip())
        return sentences

    @staticmethod
    def chunk_text_intelligently(text: str, max_tokens: int) -> List[str]:
        # Chunk text intelligently based on token count
        encoder = TokenCounter.get_encoder()
        text = PdfProcessor.preprocess(text)
        total_tokens = TokenCounter.count_tokens(text)
        if total_tokens <= max_tokens:
            return [text]

        desired_chunks = math.ceil(total_tokens / max_tokens)
        paragraphs = PdfProcessor.split_paragraphs(text)
        chunks: List[str] = []
        current = ''
        current_tokens = 0

        def flush_current():
            nonlocal current, current_tokens
            if current:
                chunks.append(current)
                current = ''
                current_tokens = 0

        for para in paragraphs:
            if PdfProcessor.is_structural(para):
                flush_current()
                chunks.append(para)
                continue

            para_tokens = TokenCounter.count_tokens(para)
            if current_tokens + para_tokens <= max_tokens:
                if current:
                    current += '\n\n' + para
                    current_tokens += para_tokens + 2
                else:
                    current = para
                    current_tokens = para_tokens
            else:
                flush_current()
                if para_tokens <= max_tokens:
                    current = para
                    current_tokens = para_tokens
                else:
                    for sent in PdfProcessor.split_sentences(para):
                        sent_tokens = TokenCounter.count_tokens(sent)
                        if sent_tokens > max_tokens:
                            chunks.append(sent)
                        elif current_tokens + sent_tokens <= max_tokens:
                            if current:
                                current += ' ' + sent
                                current_tokens += sent_tokens + 1
                            else:
                                current = sent
                                current_tokens = sent_tokens
                        else:
                            flush_current()
                            current = sent
                            current_tokens = sent_tokens
        flush_current()
        i = 0
        while len(chunks) > desired_chunks and i < len(chunks) - 1:
            tokens_i = TokenCounter.count_tokens(chunks[i])
            tokens_j = TokenCounter.count_tokens(chunks[i+1])
            if tokens_i + tokens_j <= max_tokens:
                chunks[i] = chunks[i] + '\n\n' + chunks[i+1]
                del chunks[i+1]
                i = max(i-1, 0)
            else:
                i += 1

        return chunks

class LlmClient:
    @staticmethod
    def send_prompt(full_prompt: str) -> Tuple[str, Optional[str]]:
        try:
            response = WWApiAggregator.send_prompt_to_llm(full_prompt)
            return response, None
        except Exception as e:
            return "", f"Error calling LLM API: {e}"

    @staticmethod
    def send_prompt_with_image(prompt: str, image_bytes: bytes) -> tuple[str, Optional[str]]:
        try:
            provider_name = WWSettingsManager.get_active_llm_name()
            base64_image = base64.b64encode(image_bytes).decode('utf-8')
            
            if provider_name == "LMStudio" or provider_name == "OpenAI":
                vision_content = [
                    {"type": "text", "text": prompt},
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/jpeg;base64,{base64_image}"
                        }
                    }
                ]
                conversation_history = [
                    {
                        "role": "user",
                        "content": vision_content
                    }
                ]
                response = WWApiAggregator.send_prompt_to_llm(
                    final_prompt="",
                    conversation_history=conversation_history
                )
            elif provider_name == "Anthropic":
                final_prompt = (
                    prompt +
                    "\n\n<image>\n" +
                    base64_image +
                    "\n</image>"
                )
                response = WWApiAggregator.send_prompt_to_llm(final_prompt)
            elif provider_name == "Google":
                vision_content = [
                    {"type": "text", "text": prompt},
                    {
                        "type": "image_url", 
                        "image_url": f"data:image/jpeg;base64,{base64_image}"
                    }
                ]
                conversation_history = [
                    {
                        "role": "user",
                        "content": vision_content
                    }
                ]
                response = WWApiAggregator.send_prompt_to_llm(
                    final_prompt="",
                    conversation_history=conversation_history
                )
            elif provider_name == "Ollama":
                final_prompt = f"{prompt}\n\n![image](data:image/jpeg;base64,{base64_image})"
                response = WWApiAggregator.send_prompt_to_llm(final_prompt)
            else:
                final_prompt = f"{prompt}\n\n![Image](data:image/jpeg;base64,{base64_image})"
                response = WWApiAggregator.send_prompt_to_llm(final_prompt)
            
            return response, None
        except Exception as e:
            return "", f"Error calling LLM API: {e}"

class PdfProcessingWorker(QThread):
    finished = pyqtSignal(str, list, str)
    progress = pyqtSignal(int)
    
    def __init__(self, file_path: str, pages: List[int], max_tokens: int = None):
        # Initialize with file path, sections, and optional max tokens
        super().__init__()
        self.file_path = file_path
        self.pages = pages  # Now interpreted as sections (pages, chapters, etc.) depending on processor
        self.max_tokens = max_tokens
        self.processor = DocumentProcessorFactory.get_processor(file_path)
    
    def run(self):
        # Process document using the appropriate processor
        try:
            markdown, error = self.processor.convert_to_markdown(self.file_path, self.pages)
            if error:
                self.finished.emit("", [], error)
                return
            if self.max_tokens:
                chunks = PdfProcessor.chunk_text_intelligently(markdown, self.max_tokens)
            else:
                chunks = []
            self.finished.emit(markdown, chunks, "")
        except Exception as e:
            self.finished.emit("", [], f"Error processing document: {str(e)}")

class EpubProcessingWorker(QThread):
    """
    QThread that reads an EPUB file, splits each selected chapter into
    individual <p>…</p> paragraphs, and emits:
      finished(markdown_str, list_of_paragraphs, error_str)
    """
    finished = pyqtSignal(str, list, str)

    def __init__(self, file_path: str, chapters: List[int]):
        super().__init__()
        self.file_path = file_path
        self.chapters  = chapters

    def run(self):
        try:
            # 1) Convert selected chapters into a single markdown string
            full_md, err = EpubProcessor.convert_to_markdown(
                self.file_path,
                chapters=self.chapters
            )
            if err:
                # on error, emit empty markdown + empty list + error
                self.finished.emit("", [], err)
                return

            # 2) Split on double-newline (paragraph separators)
            paras = [
                p.strip()
                for p in full_md.split("\n\n")
                if p.strip() and p.strip() != "---"
            ]

            # 3) Emit full markdown and list of small paragraphs
            self.finished.emit(full_md, paras, "")

        except Exception as e:
            self.finished.emit("", [], str(e))

class GenericProcessingWorker(QThread):
    """
    QThread that uses the DocumentProcessor API to convert any supported
    format (DOCX, TXT/MD, HTML…) into markdown, then splits that markdown
    into paragraphs/sections and emits (full_markdown, list_of_chunks, error).
    """
    finished = pyqtSignal(str, list, str)

    def __init__(self, file_path: str, sections: List[int]):
        super().__init__()
        self.file_path = file_path
        self.sections  = sections

    def run(self):
        """
        Convert DOCX, TXT/MD, or HTML into semantic blocks and emit them.
        Uses native APIs for each format to extract logical paragraphs/sections.
        """
        try:
            ext = os.path.splitext(self.file_path)[1].lower()

            # 1) DOCX: use python-docx to get real paragraphs
            if ext == '.docx':
                import docx as _docx
                doc = _docx.Document(self.file_path)
                blocks = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

            # 2) TXT or MD: split by lines, ignore empties
            elif ext in ['.txt', '.md']:
                with open(self.file_path, 'r', encoding='utf-8') as f:
                    lines = f.read().splitlines()
                blocks = [line.strip() for line in lines if line.strip()]

            # 3) HTML: extract semantic tags via BeautifulSoup
            elif ext == '.html':
                from bs4 import BeautifulSoup as _BS
                with open(self.file_path, 'r', encoding='utf-8') as f:
                    soup = _BS(f, 'html.parser')
                tags = soup.find_all(['p', 'h1', 'h2', 'h3', 'li', 'blockquote', 'div'])
                blocks = [tag.get_text(strip=True) for tag in tags if tag.get_text(strip=True)]

            # 4) Fallback: use DocumentProcessorFactory + split markdown
            else:
                from .rag_utils import DocumentProcessorFactory
                import re as _re
                processor = DocumentProcessorFactory.get_processor(self.file_path)
                full_md, err = processor.convert_to_markdown(self.file_path, chapters=self.sections)
                if err:
                    self.finished.emit("", [], err)
                    return
                blocks = [
                    chunk.strip()
                    for chunk in _re.split(r'\n{2,}', full_md)
                    if chunk.strip()
                ]

            # 5) Reassemble full markdown and emit
            full_md = "\n\n".join(blocks)
            self.finished.emit(full_md, blocks, "")

        except Exception as e:
            self.finished.emit("", [], str(e))

class SettingsManager:
    SETTINGS_FILE = "pdf_rag_settings.json"
    
    @staticmethod
    def load_settings() -> AppSettings:
        try:
            if os.path.exists(SettingsManager.SETTINGS_FILE):
                with open(SettingsManager.SETTINGS_FILE, 'r') as f:
                    data = json.load(f)
                    return AppSettings(
                        last_pdf_path_manual=data.get('last_pdf_path_manual', ""),
                        last_pdf_path_qa=data.get('last_pdf_path_qa', ""),
                        last_from_page_manual=data.get('last_from_page_manual', 0),
                        last_to_page_manual=data.get('last_to_page_manual', 0),
                        last_chunk_size=data.get('last_chunk_size', 20000),
                        default_prompt=data.get('default_prompt', "")
                    )
        except Exception as e:
            print(f"Error loading settings: {str(e)}")
        return AppSettings()
    
    @staticmethod
    def save_settings(settings: AppSettings) -> None:
        try:
            with open(SettingsManager.SETTINGS_FILE, 'w') as f:
                json.dump({
                    'last_pdf_path_manual': settings.last_pdf_path_manual,
                    'last_pdf_path_qa': settings.last_pdf_path_qa,
                    'last_from_page_manual': settings.last_from_page_manual,
                    'last_to_page_manual': settings.last_to_page_manual,
                    'last_chunk_size': settings.last_chunk_size,
                    'default_prompt': settings.default_prompt
                }, f)
        except Exception as e:
            print(f"Error saving settings: {str(e)}")

class HistoryDialog(QDialog):
    def __init__(self, parent=None, history=None):
        super().__init__(parent)
        self.history = history or []

        self.setWindowTitle("Search History")
        self.resize(600, 400)

        layout = QVBoxLayout(self)

        self.search_field = QLineEdit()
        self.search_field.setPlaceholderText("Filter history…")
        self.search_field.textChanged.connect(self.filter_history)
        layout.addWidget(self.search_field)

        self.history_list = QListWidget()
        self.history_list.setContextMenuPolicy(Qt.CustomContextMenu)
        self.history_list.customContextMenuRequested.connect(self.show_context_menu)
        self.history_list.itemDoubleClicked.connect(self.load_article)
        layout.addWidget(self.history_list)

        self.populate_history_list()

    def populate_history_list(self):
        self.history_list.clear()
        for title, _ in self.history:
            self.history_list.addItem(QListWidgetItem(title))

    def filter_history(self):
        txt = self.search_field.text().lower()
        self.history_list.clear()
        for title, _ in self.history:
            if txt in title.lower():
                self.history_list.addItem(QListWidgetItem(title))

    def show_context_menu(self, pos: QPoint):
        item = self.history_list.itemAt(pos)
        if not item:
            return
        menu = QMenu(self)
        delete = QAction("Delete", self)
        rename = QAction("Rename", self)
        menu.addAction(delete)
        menu.addAction(rename)
        delete.triggered.connect(lambda: self.delete_item(item))
        rename.triggered.connect(lambda: self.rename_item(item))
        menu.exec_(self.history_list.mapToGlobal(pos))

    def delete_item(self, item: QListWidgetItem):
        title = item.text()
        file_path = os.path.join(self.parent().history_dir, title)
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
            except Exception as e:
                QMessageBox.warning(self, "Error", f"Failed to delete file from disk: {str(e)}")
        self.history = [(t, txt) for t, txt in self.history if t != title]
        self.parent().search_history = self.history
        self.parent().save_history()
        self.populate_history_list()

    def rename_item(self, item: QListWidgetItem):
        old_title = item.text()
        name_only, ext = os.path.splitext(old_title)
        new_base, ok = QInputDialog.getText(
            self,
            "Rename Entry",
            "Enter new title (no extension):",
            text=name_only
        )
        if ok and new_base.strip():
            new_filename = new_base.strip() + ext
            for i, (t, txt) in enumerate(self.history):
                if t == old_title:
                    self.history[i] = (new_filename, txt)
                    break
            self.parent().search_history = self.history
            self.parent().save_history()
            self.populate_history_list()

    def load_article(self, item: QListWidgetItem):
        """
        Load the selected history entry (.json), display its contents in a dialog
        with search functionality, and allow saving edits back to the JSON file.
        Supports both VL-results format ("items") and QA-results format ("answer").
        """
        title = item.text()
        # Ensure the filename ends with .json
        file_name = title if title.lower().endswith(".json") else f"{title}.json"
        file_path = os.path.join(self.parent().history_dir, file_name)
        if not os.path.exists(file_path):
            QMessageBox.warning(self, "Error", f"History file not found:\n{file_path}")
            return

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to load JSON:\n{e}")
            return

        # Determine display format
        is_vl = "items" in data
        is_qa = "answer" in data
        is_manual = "chunks" in data

        if is_vl:
            # Combine all "items" entries into one block of text
            combined_text = ""
            for entry in data["items"]:
                name = entry.get("name", "")
                response = entry.get("response", "")
                combined_text += f"--- {name} ---\n{response}\n\n"
            combined_text = combined_text.rstrip()
        elif is_qa:
            question = data.get("question", "").strip()
            answer = data.get("answer", "").strip()
            combined_text = f"Question:\n{question}\n\nAnswer:\n{answer}"
        elif is_manual:
            # Combine all "chunks" entries into one block of text
            combined_text = ""
            for chunk in data["chunks"]:
                idx = chunk.get("chunk_idx", "")
                response = chunk.get("response", "").strip()
                combined_text += f"--- Chunk {idx} ---\n{response}\n\n"
            combined_text = combined_text.rstrip()
        else:
            combined_text = json.dumps(data, indent=2, ensure_ascii=False)

        dlg = QDialog(self)
        dlg.setWindowTitle(f"History: {file_name}")
        dlg.resize(800, 600)
        main_layout = QVBoxLayout(dlg)

        header_label = QLabel(f"<b>{file_name}</b>")
        main_layout.addWidget(header_label)

        # Search bar (hidden initially)
        search_container = QWidget()
        search_layout = QHBoxLayout(search_container)
        search_layout.setContentsMargins(0, 0, 0, 0)

        search_field = QLineEdit()
        search_field.setPlaceholderText("Search text…")
        search_layout.addWidget(search_field)

        case_checkbox = QCheckBox("Case sensitive")
        search_layout.addWidget(case_checkbox)

        whole_words_checkbox = QCheckBox("Whole words")
        search_layout.addWidget(whole_words_checkbox)

        prev_btn = QPushButton("Previous")
        search_layout.addWidget(prev_btn)

        next_btn = QPushButton("Next")
        search_layout.addWidget(next_btn)

        close_search_btn = QPushButton("×")
        close_search_btn.setFixedSize(25, 25)
        close_search_btn.setToolTip("Close search")
        search_layout.addWidget(close_search_btn)

        search_container.setVisible(False)
        main_layout.addWidget(search_container)

        # Text editor for combined_text
        editor = QTextEdit()
        editor.setPlainText(combined_text)
        main_layout.addWidget(editor)

        # Buttons: Save and Close
        button_layout = QHBoxLayout()
        save_btn = QPushButton("Save")
        save_btn.setEnabled(False)
        button_layout.addWidget(save_btn)
        close_btn = QPushButton("Close")
        button_layout.addWidget(close_btn)
        main_layout.addLayout(button_layout)

        original_text = combined_text

        def toggle_search_bar():
            search_container.setVisible(not search_container.isVisible())
            if search_container.isVisible():
                search_field.setFocus()
                search_field.selectAll()

        def find_text(direction=1):
            search_text = search_field.text()
            if not search_text:
                return
            flags = QTextDocument.FindFlags()
            if case_checkbox.isChecked():
                flags |= QTextDocument.FindCaseSensitively
            if whole_words_checkbox.isChecked():
                flags |= QTextDocument.FindWholeWords
            if direction < 0:
                flags |= QTextDocument.FindBackward

            cursor = editor.textCursor()
            if direction < 0 and cursor.hasSelection():
                pos = cursor.selectionStart()
                cursor.setPosition(pos)
                editor.setTextCursor(cursor)

            found = editor.find(search_text, flags)
            if not found:
                temp_cursor = editor.textCursor()
                cursor = editor.textCursor()
                if direction > 0:
                    cursor.movePosition(QTextCursor.Start)
                else:
                    cursor.movePosition(QTextCursor.End)
                editor.setTextCursor(cursor)
                found = editor.find(search_text, flags)
                if not found:
                    editor.setTextCursor(temp_cursor)
                    QMessageBox.information(dlg, "Search Result", f"No matches for '{search_text}'")

        def on_text_changed():
            save_btn.setEnabled(editor.toPlainText() != original_text)

        def save_changes():
            new_text = editor.toPlainText()
            try:
                if is_vl:
                    # Preserve entire edited text as a single "Combined" entry
                    data["items"] = [{"name": "Combined", "response": new_text}]
                    # Remove QA keys if present
                    data.pop("question", None)
                    data.pop("answer", None)
                elif is_qa:
                    # Expect format "Question:\n...\n\nAnswer:\n..."
                    parts = new_text.split("\n\nAnswer:\n", 1)
                    if len(parts) == 2 and parts[0].startswith("Question:\n"):
                        question_text = parts[0][len("Question:\n"):].strip()
                        answer_text = parts[1].strip()
                        data["question"] = question_text
                        data["answer"] = answer_text
                    data.pop("items", None)
                elif is_manual:
                    # Split edited text back into chunks if needed (overwrite with single "Combined" chunk)
                    data["chunks"] = [{"chunk_idx": 0, "response": new_text}]
                    data.pop("question", None)
                    data.pop("answer", None)
                    data.pop("items", None)
                else:
                    # Overwrite entire JSON with valid JSON from editor
                    parsed = json.loads(new_text)
                    data.clear()
                    data.update(parsed)

                with open(file_path, 'w', encoding='utf-8') as f_out:
                    json.dump(data, f_out, indent=2, ensure_ascii=False)

                # Update in-memory history list
                parent_history = self.parent().search_history
                for i, (t, txt) in enumerate(parent_history):
                    if t == title:
                        parent_history[i] = (title, new_text)
                        break
                self.parent().search_history = parent_history
                self.parent().save_history()

                nonlocal original_text
                original_text = new_text
                save_btn.setEnabled(False)
                QMessageBox.information(dlg, "Saved", "Changes have been saved to JSON.")
            except Exception as e:
                QMessageBox.warning(dlg, "Error", f"Failed to save JSON:\n{e}")

        # Connect signals
        editor.textChanged.connect(on_text_changed)
        save_btn.clicked.connect(save_changes)
        close_btn.clicked.connect(dlg.close)

        shortcut_find = QShortcut(QKeySequence("Ctrl+F"), dlg)
        shortcut_find.activated.connect(toggle_search_bar)
        close_search_btn.clicked.connect(toggle_search_bar)
        search_field.returnPressed.connect(lambda: find_text(1))
        next_btn.clicked.connect(lambda: find_text(1))
        prev_btn.clicked.connect(lambda: find_text(-1))
        shortcut_next = QShortcut(QKeySequence("F3"), dlg)
        shortcut_next.activated.connect(lambda: find_text(1))
        shortcut_prev = QShortcut(QKeySequence("Shift+F3"), dlg)
        shortcut_prev.activated.connect(lambda: find_text(-1))
        shortcut_close = QShortcut(QKeySequence("Escape"), dlg)
        shortcut_close.activated.connect(lambda: 
            search_container.setVisible(False) if search_container.isVisible() else None
        )

        dlg.exec_()